"""
Test utilities - create sample DOCX files for testing.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def create_test_template(path: Path) -> None:
    """Create a sample template DOCX with professional styles."""
    doc = Document()

    # Set document margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Create custom styles
    # Title style
    title_style = doc.styles.add_style('OfferTitle', 1)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 153)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(18)

    # Heading 1 style
    h1_style = doc.styles.add_style('OfferHeading1', 1)
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 153)
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(6)

    # Heading 2 style
    h2_style = doc.styles.add_style('OfferHeading2', 1)
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(13)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(51, 102, 204)
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(6)

    # Body style
    body_style = doc.styles.add_style('OfferBody', 1)
    body_style.font.name = 'Calibri'
    body_style.font.size = Pt(11)
    body_style.paragraph_format.line_spacing = 1.15
    body_style.paragraph_format.space_after = Pt(6)

    # Quote style
    quote_style = doc.styles.add_style('OfferQuote', 1)
    quote_style.font.name = 'Calibri'
    quote_style.font.size = Pt(11)
    quote_style.font.italic = True
    quote_style.font.color.rgb = RGBColor(100, 100, 100)
    quote_style.paragraph_format.left_indent = Inches(0.5)
    quote_style.paragraph_format.right_indent = Inches(0.5)
    quote_style.paragraph_format.space_before = Pt(6)
    quote_style.paragraph_format.space_after = Pt(6)

    # Add some placeholder content
    doc.add_paragraph('OFERTA HANDLOWA', style='OfferTitle')
    doc.add_paragraph('Dla: {{nazwa_firmy}}', style='OfferBody')
    doc.add_paragraph('Data: {{data}}', style='OfferBody')

    doc.add_paragraph('PRZEDMIOT OFERTY', style='OfferHeading1')
    doc.add_paragraph('W ramach niniejszej oferty przedstawiamy szczegółowe informacje...', style='OfferBody')

    doc.add_paragraph('WARUNKI WSPÓŁPRACY', style='OfferHeading2')
    doc.add_paragraph('Warunki płatności: przelew 14 dni', style='OfferBody')

    doc.add_paragraph('Podpisano przez:', style='OfferBody')
    doc.add_paragraph('The quick brown fox', style='OfferQuote')

    doc.save(str(path))
    print(f"Created template: {path}")


def create_test_content(path: Path) -> None:
    """Create a sample content DOCX with raw text."""
    doc = Document()

    # Add content without proper styling (simulating raw input)
    doc.add_paragraph('Oferta na system CRM')

    p = doc.add_paragraph('Dla: ACME Sp. z o.o.')
    p.runs[0].bold = True

    p = doc.add_paragraph('23 stycznia 2025')
    p.runs[0].italic = True

    doc.add_paragraph('')
    doc.add_paragraph('Oferujemy wdrożenie systemu CRM dostosowanego do potrzeb Państwa firmy.')

    doc.add_paragraph('ZAKRES PRAC')
    doc.add_paragraph('1. Analiza wymagań biznesowych')
    doc.add_paragraph('2. Konfiguracja systemu')
    doc.add_paragraph('3. Szkolenia dla użytkowników')

    doc.add_paragraph('Koszty')
    doc.add_paragraph('Całkowity koszt projektu: 45 000 PLN')
    doc.add_paragraph('Termin realizacji: 3 miesiące')

    doc.add_paragraph('')
    doc.add_paragraph('Z poważaniem,')
    doc.add_paragraph('Jan Kowalski')

    doc.save(str(path))
    print(f"Created content: {path}")


def create_cv_template(path: Path) -> None:
    """Create a CV template DOCX."""
    doc = Document()

    # Set narrow margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Personal info style
    name_style = doc.styles.add_style('CVName', 1)
    name_style.font.name = 'Georgia'
    name_style.font.size = Pt(22)
    name_style.font.bold = True
    name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section header style
    section_style = doc.styles.add_style('CVSection', 1)
    section_style.font.name = 'Georgia'
    section_style.font.size = Pt(14)
    section_style.font.bold = True
    section_style.font.color.rgb = RGBColor(0, 102, 102)
    section_style.paragraph_format.space_before = Pt(12)
    section_style.paragraph_format.space_after = Pt(4)
    section_style.paragraph_format.bottom_border = True  # Will add border via XML

    # Body style
    body_style = doc.styles.add_style('CVBody', 1)
    body_style.font.name = 'Calibri'
    body_style.font.size = Pt(10.5)
    body_style.paragraph_format.space_after = Pt(3)

    # Add placeholder content
    doc.add_paragraph('IMIĘ I NAZWISKO', style='CVName')
    doc.add_paragraph('email@example.com | +48 123 456 789 | Warszawa', style='CVBody')

    doc.add_paragraph('DOŚWIADCZENIE', style='CVSection')
    doc.add_paragraph('Stanowisko | Firma | 2020-2024', style='CVBody')
    doc.add_paragraph('Opis obowiązków i osiągnięć', style='CVBody')

    doc.add_paragraph('WYKSZTAŁCENIE', style='CVSection')
    doc.add_paragraph('Tytuł | Uczelnia | 2015-2019', style='CVBody')

    doc.save(str(path))
    print(f"Created CV template: {path}")


def create_cv_content(path: Path) -> None:
    """Create a sample CV content DOCX."""
    doc = Document()

    doc.add_paragraph('Paweł Nowak')
    doc.add_paragraph('pawel.nowak@email.com | +48 600 700 800 | Kraków')

    doc.add_paragraph('')
    doc.add_paragraph('DOŚWIADCZENIE ZAWODOWE')

    p = doc.add_paragraph('Senior Python Developer')
    p.runs[0].bold = True
    doc.add_paragraph('Nous Research | 2022-obecnie')
    doc.add_paragraph('Budowa autonomicznych agentów AI')
    doc.add_paragraph('')

    p = doc.add_paragraph('Full Stack Developer')
    p.runs[0].bold = True
    doc.add_paragraph('TechCorp | 2019-2022')
    doc.add_paragraph('Rozwój platformy SaaS')

    doc.add_paragraph('')
    doc.add_paragraph('WYKSZTAŁCENIE')
    doc.add_paragraph('Magister Informatyki | AGH | 2017-2019')
    doc.add_paragraph('Inżynier Informatyki | AGH | 2013-2017')

    doc.add_paragraph('')
    doc.add_paragraph('UMIEJĘTNOŚCI')
    doc.add_paragraph('Python, TypeScript, React, Docker, Kubernetes')

    doc.save(str(path))
    print(f"Created CV content: {path}")


if __name__ == '__main__':
    import sys
    target_dir = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('/Users/pawel/docx-formatter/tests/fixtures')
    target_dir.mkdir(parents=True, exist_ok=True)

    create_test_template(target_dir / 'template_offer.docx')
    create_test_content(target_dir / 'content_offer.docx')
    create_cv_template(target_dir / 'template_cv.docx')
    create_cv_content(target_dir / 'content_cv.docx')

    print("\nDone! Created 4 test files.")
