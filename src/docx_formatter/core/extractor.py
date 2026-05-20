"""
DOCX Extractor - parses both template and content documents into profiles.
"""

import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Optional, Dict, List, Any, Tuple
import logging

# Try importing python-docx
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .types import (
    FontStyle, ParagraphSpacing, ParagraphStyle, CharacterStyle,
    TableStyle, DocumentDefaults, Theme, HeaderFooter, SectionBreak,
    TemplateProfile, ContentProfile, ParagraphContent, TableContent,
    SemanticRole, TemplateType, ParagraphAlignment, ParagraphIndentation,
    ParagraphBorder
)

logger = logging.getLogger(__name__)

# Namespaces OOXML
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
}


def _get_text(elem: ET.Element) -> str:
    """Extract all text from an XML element recursively."""
    text = []
    if elem.text:
        text.append(elem.text)
    for child in elem:
        text.append(_get_text(child))
        if child.tail:
            text.append(child.tail)
    return ''.join(text)


def _parse_color(color_elem: Optional[ET.Element]) -> Optional[str]:
    """Parse color from w:color element. Returns hex string like 'FF0000'."""
    if color_elem is None:
        return None
    val = color_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
    if val and val != 'auto':
        return val.upper()
    # Try theme color
    theme_color = color_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeColor')
    if theme_color:
        return f"theme:{theme_color}"
    return None


def _parse_font_size(sz_elem: Optional[ET.Element]) -> Optional[float]:
    """Parse font size in half-points, return points."""
    if sz_elem is None:
        return None
    val = sz_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
    if val:
        try:
            return int(val) / 2.0
        except (ValueError, TypeError):
            pass
    return None


class DOCXExtractor:
    """Parses DOCX files into structured profiles."""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    def extract_template_profile(self, docx_path: str) -> TemplateProfile:
        """
        Extract template profile from Document A (template).
        
        Returns all style definitions, document settings, theme, headers/footers.
        """
        doc = Document(docx_path)
        profile = TemplateProfile()
        
        # Extract paragraph styles
        profile.paragraph_styles = self._extract_paragraph_styles(doc)
        logger.info(f"Extracted {len(profile.paragraph_styles)} paragraph styles")
        
        # Extract document defaults (margins, page size)
        profile.document_defaults = self._extract_document_defaults(doc)
        
        # Try to read styles.xml directly for full style definitions
        self._extract_styles_xml(docx_path, profile)
        
        # Detect template type
        profile.template_type = self._detect_template_type(doc, profile)
        
        return profile
    
    def extract_content_profile(self, docx_path: str) -> ContentProfile:
        """
        Extract content profile from Document B (raw content).
        
        Returns paragraphs, tables, images with their current formatting.
        """
        doc = Document(docx_path)
        profile = ContentProfile()
        
        # Extract paragraphs
        for para in doc.paragraphs:
            para_content = self._extract_paragraph(para)
            if para_content.text.strip():
                profile.paragraphs.append(para_content)
                profile.word_count += len(para_content.text.split())
        
        # Extract tables
        for table in doc.tables:
            table_content = self._extract_table(table)
            if table_content.num_rows > 0:
                profile.tables.append(table_content)
        
        # Build structure tree
        profile.structure_tree = self._build_structure_tree(profile.paragraphs)
        
        # Guess document title from first heading or paragraph
        for para in profile.paragraphs:
            if para.estimated_role in (SemanticRole.TITLE, SemanticRole.HEADING_1):
                profile.title = para.text[:200]
                break
        if not profile.title and profile.paragraphs:
            profile.title = profile.paragraphs[0].text[:200]
        
        logger.info(f"Extracted {len(profile.paragraphs)} paragraphs, {len(profile.tables)} tables")
        return profile
    
    def _extract_paragraph_styles(self, doc) -> Dict[str, ParagraphStyle]:
        """Extract all paragraph styles from document."""
        styles = {}
        for style in doc.styles:
            if style.type is None:
                continue
            try:
                style_type = style.type
            except Exception:
                continue
            
            if style_type != 1:  # WD_STYLE_TYPE.PARAGRAPH = 1
                continue
            
            style_id = style.style_id
            font = FontStyle()
            
            # Font properties
            try:
                if style.font.name:
                    font.name = style.font.name
                if style.font.size:
                    font.size_pt = style.font.size.pt
                font.bold = style.font.bold
                font.italic = style.font.italic
                if style.font.color and style.font.color.rgb:
                    font.color = str(style.font.color.rgb)
            except Exception as e:
                logger.debug(f"Error reading font for style {style_id}: {e}")
            
            # Paragraph format
            alignment = ParagraphAlignment()
            spacing = ParagraphSpacing()
            indentation = ParagraphIndentation()
            
            try:
                pf = style.paragraph_format
                if pf.alignment is not None:
                    alignment.alignment = str(pf.alignment)
                if pf.space_before:
                    spacing.before_pt = pf.space_before.pt
                if pf.space_after:
                    spacing.after_pt = pf.space_after.pt
                if pf.line_spacing:
                    spacing.line_spacing = float(pf.line_spacing)
                if pf.left_indent:
                    indentation.left_inches = pf.left_indent.inches
                if pf.right_indent:
                    indentation.right_inches = pf.right_indent.inches
                if pf.first_line_indent:
                    indentation.first_line_inches = pf.first_line_indent.inches
            except Exception as e:
                logger.debug(f"Error reading paragraph format for style {style_id}: {e}")
            
            ps = ParagraphStyle(
                style_id=style_id,
                name=style.name,
                based_on=style.base_style.style_id if style.base_style else None,
                font=font,
                alignment=alignment,
                spacing=spacing,
                indentation=indentation,
                is_built_in=style.builtin,
            )
            styles[style_id] = ps
            
        return styles
    
    def _extract_document_defaults(self, doc) -> DocumentDefaults:
        """Extract document-level defaults."""
        defaults = DocumentDefaults()
        try:
            section = doc.sections[0]
            defaults.margins = {
                'top': section.top_margin.inches,
                'bottom': section.bottom_margin.inches,
                'left': section.left_margin.inches,
                'right': section.right_margin.inches,
            }
            defaults.page_width_inches = section.page_width.inches
            defaults.page_height_inches = section.page_height.inches
            defaults.orientation = 'landscape' if section.orientation else 'portrait'
        except Exception as e:
            logger.warning(f"Error extracting document defaults: {e}")
        return defaults
    
    def _extract_styles_xml(self, docx_path: str, profile: TemplateProfile) -> None:
        """
        Read styles.xml directly for full style definitions including outline levels.
        This provides more complete data than python-docx API.
        """
        try:
            with zipfile.ZipFile(docx_path, 'r') as zf:
                if 'word/styles.xml' not in zf.namelist():
                    return
                
                styles_xml = zf.read('word/styles.xml')
                root = ET.fromstring(styles_xml)
                ns = NAMESPACES['w']
                
                for style_elem in root.findall(f'.//{{{ns}}}style'):
                    style_id = style_elem.get(f'{{{ns}}}styleId')
                    style_type = style_elem.get(f'{{{ns}}}type')
                    
                    if style_type != 'paragraph':
                        continue
                    
                    # Find or create style
                    if style_id in profile.paragraph_styles:
                        ps = profile.paragraph_styles[style_id]
                    else:
                        ps = ParagraphStyle(style_id=style_id, name=style_id)
                        profile.paragraph_styles[style_id] = ps
                    
                    # Extract name
                    name_elem = style_elem.find(f'.//{{{ns}}}name')
                    if name_elem is not None:
                        ps.name = name_elem.get(f'{{{ns}}}val') or ps.name
                    
                    # Extract basedOn
                    based_elem = style_elem.find(f'{{{ns}}}basedOn')
                    if based_elem is not None:
                        ps.based_on = based_elem.get(f'{{{ns}}}val')
                    
                    # Outline level
                    outline_elem = style_elem.find(f'.//{{{ns}}}outlineLvl')
                    if outline_elem is not None:
                        val = outline_elem.get(f'{{{ns}}}val')
                        if val is not None:
                            ps.outline_level = int(val)
                            # Auto-assign semantic role for headings
                            if ps.outline_level is not None:
                                outline_role_map = {
                                    0: SemanticRole.HEADING_1,
                                    1: SemanticRole.HEADING_2,
                                    2: SemanticRole.HEADING_3,
                                    3: SemanticRole.HEADING_4,
                                }
                                ps.semantic_role = outline_role_map.get(ps.outline_level)
                                ps.role_confidence = 0.9
                    
                    # Extract font properties from rPr if not already set
                    rpr = style_elem.find(f'.//{{{ns}}}rPr')
                    if rpr is not None:
                        if ps.font.size_pt is None:
                            sz = rpr.find(f'{{{ns}}}sz')
                            ps.font.size_pt = _parse_font_size(sz)
                        if ps.font.color is None:
                            color = rpr.find(f'{{{ns}}}color')
                            ps.font.color = _parse_color(color)
                        if ps.font.name is None:
                            rfonts = rpr.find(f'{{{ns}}}rFonts')
                            if rfonts is not None:
                                ps.font.name = rfonts.get(f'{{{ns}}}ascii')
                    
                    # Paragraph properties
                    ppr = style_elem.find(f'{{{ns}}}pPr')
                    if ppr is not None:
                        # Spacing
                        spacing_elem = ppr.find(f'{{{ns}}}spacing')
                        if spacing_elem is not None:
                            before = spacing_elem.get(f'{{{ns}}}before')
                            after = spacing_elem.get(f'{{{ns}}}after')
                            if before:
                                ps.spacing.before_pt = int(before) / 20.0  # twips to pt
                            if after:
                                ps.spacing.after_pt = int(after) / 20.0
                            line = spacing_elem.get(f'{{{ns}}}line')
                            if line:
                                line_val = int(line)
                                line_rule = spacing_elem.get(f'{{{ns}}}lineRule')
                                if line_rule == 'auto':
                                    ps.spacing.line_spacing = line_val / 240.0
                                else:
                                    ps.spacing.line_spacing = line_val / 20.0
                        
                        # Alignment
                        jc = ppr.find(f'{{{ns}}}jc')
                        if jc is not None:
                            ps.alignment.alignment = jc.get(f'{{{ns}}}val')
        
        except Exception as e:
            logger.warning(f"Error reading styles.xml: {e}")
    
    def _extract_paragraph(self, para) -> ParagraphContent:
        """Extract paragraph content and metadata."""
        text = para.text or ''
        
        # Style name
        style_name = None
        try:
            if para.style:
                style_name = para.style.style_id
        except Exception:
            pass
        
        # Runs with formatting
        runs = []
        has_bold = False
        has_italic = False
        has_underline = False
        max_font_size = None
        
        for run in para.runs:
            run_text = run.text or ''
            if not run_text:
                continue
            
            run_info = {'text': run_text}
            try:
                if run.bold:
                    run_info['bold'] = True
                    has_bold = True
                if run.italic:
                    run_info['italic'] = True
                    has_italic = True
                if run.underline:
                    run_info['underline'] = True
                    has_underline = True
                if run.font.size:
                    size = run.font.size.pt
                    run_info['size'] = size
                    if max_font_size is None or size > max_font_size:
                        max_font_size = size
                if run.font.color and run.font.color.rgb:
                    run_info['color'] = str(run.font.color.rgb)
            except Exception:
                pass
            
            runs.append(run_info)
        
        # Check for list formatting
        is_list = False
        list_level = 0
        left_indent = None
        try:
            pf = para.paragraph_format
            if pf.left_indent is not None:
                left_indent = pf.left_indent.inches
        except Exception:
            pass
            left_indent = None
        
        # Estimated semantic role from content
        estimated_role = self._estimate_role(text, has_bold, max_font_size, style_name)
        
        return ParagraphContent(
            text=text,
            style_name=style_name,
            runs=runs,
            is_list_item=is_list,
            list_level=list_level,
            indentation_left_inches=left_indent,
            has_bold=has_bold,
            has_italic=has_italic,
            has_underline=has_underline,
            max_font_size=max_font_size,
            estimated_role=estimated_role,
        )
    
    def _estimate_role(self, text: str, has_bold: bool, max_font_size: Optional[float],
                       style_name: Optional[str]) -> Optional[SemanticRole]:
        """Estimate semantic role from content heuristics."""
        text = text.strip()
        if not text:
            return None
        
        # Check style name hints
        if style_name:
            s = style_name.lower()
            role_map = {
                'heading 1': SemanticRole.HEADING_1,
                'heading 2': SemanticRole.HEADING_2,
                'heading 3': SemanticRole.HEADING_3,
                'heading 4': SemanticRole.HEADING_4,
                'title': SemanticRole.TITLE,
                'subtitle': SemanticRole.SUBTITLE,
                'quote': SemanticRole.QUOTE,
                'caption': SemanticRole.CAPTION,
                'list': SemanticRole.LIST_BULLET,
                'bullet': SemanticRole.LIST_BULLET,
                'numbered': SemanticRole.LIST_NUMBER,
                'footnote': SemanticRole.FOOTER,
            }
            for key, role in role_map.items():
                if key in s:
                    return role
        
        # Heuristic based on formatting
        if max_font_size and max_font_size >= 20:
            if has_bold:
                return SemanticRole.TITLE
        
        if max_font_size and max_font_size >= 16:
            if has_bold:
                return SemanticRole.HEADING_1
        
        if text.startswith('●') or text.startswith('•') or text.startswith('○'):
            return SemanticRole.LIST_BULLET
        
        # Date patterns
        import re
        date_patterns = [
            r'\d{2}[./]\d{2}[./]\d{4}',
            r'\d{4}-\d{2}-\d{2}',
        ]
        for pattern in date_patterns:
            if re.search(pattern, text):
                return SemanticRole.DATE_FIELD
        
        # Amount patterns
        amount_pattern = r'\d{1,3}(?:\s?\d{3})*(?:[.,]\d{2})?\s*(?:PLN|USD|EUR|GBP|\$|€|£)'
        if re.search(amount_pattern, text):
            return SemanticRole.AMOUNT_FIELD
        
        return SemanticRole.BODY_TEXT
    
    def _extract_table(self, table) -> TableContent:
        """Extract table content."""
        rows = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            rows.append(row_data)
        
        style_name = None
        try:
            if table.style:
                style_name = table.style.style_id
        except Exception:
            pass
        
        return TableContent(
            rows=rows,
            style_name=style_name,
            num_rows=len(rows),
            num_cols=len(rows[0]) if rows else 0,
        )
    
    def _build_structure_tree(self, paragraphs: List[ParagraphContent]) -> Dict[str, Any]:
        """Build hierarchical structure tree from paragraphs."""
        tree = {'type': 'document', 'children': []}
        current_section = tree
        
        for para in paragraphs:
            role = para.estimated_role or SemanticRole.BODY_TEXT
            
            if role in (SemanticRole.TITLE, SemanticRole.HEADING_1):
                section = {
                    'type': 'section',
                    'heading': para.text,
                    'role': role.value,
                    'children': []
                }
                tree['children'].append(section)
                current_section = section
            elif role in (SemanticRole.HEADING_2, SemanticRole.HEADING_3):
                subsection = {
                    'type': 'subsection',
                    'heading': para.text,
                    'role': role.value,
                    'children': []
                }
                current_section['children'].append(subsection)
                current_section = subsection
            else:
                current_section.setdefault('children', []).append({
                    'type': 'paragraph',
                    'text': para.text[:100],
                    'role': role.value,
                })
        
        return tree
    
    def _detect_template_type(self, doc, profile: TemplateProfile) -> TemplateType:
        """Detect what kind of template this is."""
        total_paragraphs = len(doc.paragraphs)
        text_paragraphs = sum(1 for p in doc.paragraphs if p.text.strip())
        
        # If mostly empty or only "Lorem ipsum" content -> style definitions
        if text_paragraphs <= 3:
            return TemplateType.STYLE_DEFINITIONS
        
        # Check if content looks like instructions
        instruction_keywords = ['format', 'style', 'template', 'document', 'heading', 'font']
        full_text = ' '.join(p.text.lower() for p in doc.paragraphs)
        instruction_score = sum(1 for kw in instruction_keywords if kw in full_text)
        
        if instruction_score >= 5:
            return TemplateType.TEXT_INSTRUCTIONS
        
        return TemplateType.EXAMPLE_DOCUMENT


__all__ = ['DOCXExtractor']
