"""
Document Assembler - builds final DOCX output from template + content + matches.

Strategy when template_docx_path is provided:
    1. Open the template DOCX as the base (preserves ZIP, media, headers, footers).
    2. Categorise each section in the template as:
       - structural  → TOC, front matter, cover design  (preserved but text cleared)
       - content     → body text chapters, executive summaries, appendices (filled with content)
    3. In content-bearing sections, replace template paragraphs in order using
       semantic slot matching (Heading 1 ↔ Heading 1, Body text ↔ Body text).
    4. Append unmatched content paras, and clear unused body-text placeholders.
"""

import copy
import logging
from typing import Dict, List, Optional, Any
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .types import (
    TemplateProfile, ContentProfile, ParagraphContent,
    ParagraphStyle, StyleMatch, ProcessingResult,
    SemanticRole, FontStyle
)

logger = logging.getLogger(__name__)


class DocumentAssembler:
    """Assembles the final DOCX document from template + content + style matches."""

    def __init__(self, preserve_template_content: bool = False):
        self.preserve_template_content = preserve_template_content

    # ── Public API ───────────────────────────────────────

    def assemble(
        self,
        template: TemplateProfile,
        content: ContentProfile,
        style_matches: List[StyleMatch],
        output_path: Optional[str] = None,
        template_docx_path: Optional[str] = None,
    ) -> ProcessingResult:
        result = ProcessingResult()
        try:
            if template_docx_path:
                doc = Document(template_docx_path)
                self._replace_in_template(doc, template, content, style_matches)
            else:
                doc = Document()
                self._build_fresh(doc, template, content, style_matches)

            if output_path:
                doc.save(output_path)
                result.output_path = output_path
                logger.info(f"Document saved to {output_path}")

            result.matched_styles = style_matches
            result.success = True
            result.warnings = []

        except Exception as e:
            logger.error(f"Assembly failed: {e}", exc_info=True)
            result.warnings.append(str(e))
            result.success = False
        return result

    # ── Main replacement logic ───────────────────────────

    def _replace_in_template(
        self,
        doc: Document,
        template: TemplateProfile,
        content: ContentProfile,
        style_matches: List[StyleMatch],
    ) -> None:
        """Replace template paragraph text based on section type.

        Strategy:
        1. Split content into two queues: headings and body text (preserves order).
        2. In content sections, fill heading slots from heading queue and body slots
           from body queue. This prevents headings/body text from mixing.
        3. Structural sections (TOC, frontmatter) are left untouched.
        4. Overflow from both queues is inserted into the last content section.
        """
        style_map = {m.source_style_id: m.target_style_id for m in style_matches}
        sections = self._get_section_boundaries(doc)

        # Split content into ordered queues, respecting style mappings
        heading_style_ids = set()
        for sid, tstyle in template.paragraph_styles.items():
            if tstyle.semantic_role in (SemanticRole.TITLE, SemanticRole.HEADING_1,
                                        SemanticRole.HEADING_2, SemanticRole.HEADING_3,
                                        SemanticRole.HEADING_4):
                heading_style_ids.add(sid)

        content_headings = []
        content_body = []
        for p in content.paragraphs:
            # Is this content's style mapped to a heading?
            mapped_to_heading = (p.style_name and p.style_name in style_map
                                 and style_map[p.style_name] in heading_style_ids)
            # Is this explicitly a heading style?
            style_is_heading = p.style_name and "heading" in p.style_name.lower()
            # Heuristic role
            role = self._resolve_content_role(p)

            if mapped_to_heading or style_is_heading or role in (
                SemanticRole.TITLE, SemanticRole.HEADING_1,
                SemanticRole.HEADING_2, SemanticRole.HEADING_3,
                SemanticRole.HEADING_4,
            ):
                content_headings.append(p)
            else:
                content_body.append(p)

        h_idx = 0
        b_idx = 0
        last_content_section_end = None

        # Pass 1: identify last content section
        for sec_idx, (start, end) in enumerate(sections):
            if self._classify_section(doc, start, end) == "content":
                last_content_section_end = end

        # Pass 2: fill slots from appropriate queues
        for sec_idx, (start, end) in enumerate(sections):
            sec_role = self._classify_section(doc, start, end)

            if sec_role == "content":
                for t_idx in range(start, end + 1):
                    slot = self._slot_type(doc.paragraphs[t_idx])
                    if slot is None:
                        self._clear_para_runs(doc.paragraphs[t_idx]._element)
                        continue

                    if slot == "heading" and h_idx < len(content_headings):
                        self._replace_para_text_keep_structure(
                            doc.paragraphs[t_idx]._element,
                            content_headings[h_idx],
                            template, style_map,
                        )
                        h_idx += 1
                    elif slot == "body" and b_idx < len(content_body):
                        self._replace_para_text_keep_structure(
                            doc.paragraphs[t_idx]._element,
                            content_body[b_idx],
                            template, style_map,
                        )
                        b_idx += 1
                    else:
                        self._clear_para_runs(doc.paragraphs[t_idx]._element)

            elif sec_role == "cover":
                cover_filled = False
                for t_idx in range(start, end + 1):
                    slot = self._slot_type(doc.paragraphs[t_idx])
                    if slot == "heading" and not cover_filled and h_idx < len(content_headings):
                        self._replace_para_text_keep_structure(
                            doc.paragraphs[t_idx]._element,
                            content_headings[h_idx],
                            template, style_map,
                        )
                        h_idx += 1
                        cover_filled = True
                    elif slot in ("heading", "body"):
                        self._clear_para_runs(doc.paragraphs[t_idx]._element)

            elif sec_role in ("toc", "frontmatter"):
                if not self.preserve_template_content:
                    for t_idx in range(start, end + 1):
                        if self._slot_type(doc.paragraphs[t_idx]) == "body":
                            self._clear_para_runs(doc.paragraphs[t_idx]._element)

        # Pass 3: merge remaining headings + body text in original order and insert
        remaining = self._merge_remaining(content.paragraphs, content_headings[h_idx:], content_body[b_idx:])
        if remaining and last_content_section_end is not None:
            self._insert_overflow_before_section_break(
                doc, last_content_section_end, remaining, template, style_map,
            )

    def _merge_remaining(self, all_paras: List[ParagraphContent],
                         remaining_headings: List[ParagraphContent],
                         remaining_body: List[ParagraphContent]) -> List[ParagraphContent]:
        """Reconstruct remaining paragraphs in original document order."""
        remaining_set = set(id(p) for p in remaining_headings + remaining_body)
        return [p for p in all_paras if id(p) in remaining_set]

    def _insert_overflow_before_section_break(
        self,
        doc: Document,
        section_end_para_idx: int,
        remaining: List[ParagraphContent],
        template: TemplateProfile,
        style_map: Dict[str, str],
    ) -> None:
        """Insert remaining content paragraphs before the section-break paragraph."""
        body_elem = doc.element.body
        # Find the section-break paragraph element
        sect_break_elem = doc.paragraphs[section_end_para_idx]._element
        # Find its index in the body
        try:
            insert_idx = list(body_elem).index(sect_break_elem)
        except ValueError:
            insert_idx = len(list(body_elem))

        for para in remaining:
            target_style_id = self._resolve_target_style_for_append(para, style_map, template)
            p_elem = self._new_paragraph_element(para, target_style_id)
            body_elem.insert(insert_idx, p_elem)
            insert_idx += 1

    def _resolve_content_role(self, para: ParagraphContent) -> SemanticRole:
        """Determine semantic role, prioritising explicit style names over heuristics."""
        # If the paragraph has a heading style, trust it
        style = (para.style_name or "").lower()
        if "heading" in style or "title" in style:
            if "heading 1" in style or "heading1" in style:
                return SemanticRole.HEADING_1
            if "heading 2" in style or "heading2" in style:
                return SemanticRole.HEADING_2
            if "heading 3" in style or "heading3" in style:
                return SemanticRole.HEADING_3
            if "heading 4" in style or "heading4" in style:
                return SemanticRole.HEADING_4
            if "title" in style:
                return SemanticRole.TITLE
            return SemanticRole.HEADING_1

        if para.estimated_role and para.estimated_role != SemanticRole.UNKNOWN:
            return para.estimated_role

        text = para.text.strip()
        if not text:
            return SemanticRole.UNKNOWN
        if len(text) < 100:
            if para.max_font_size and para.max_font_size >= 20:
                return SemanticRole.TITLE
            if para.max_font_size and para.max_font_size >= 16:
                return SemanticRole.HEADING_1
            if para.max_font_size and para.max_font_size >= 12:
                return SemanticRole.HEADING_2
            if para.has_bold and len(text) < 80:
                return SemanticRole.HEADING_2
        if para.is_list_item or text.startswith(('•', '●', '○', '▪', '-')):
            return SemanticRole.LIST_BULLET
        if text[:3].strip().endswith('.') and text[:2].strip()[0].isdigit():
            return SemanticRole.LIST_NUMBER
        return SemanticRole.BODY_TEXT

    def _get_section_boundaries(self, doc: Document) -> List[tuple]:
        """Return list of (start_para_index, end_para_index) for each section."""
        breaks = [
            i for i, p in enumerate(doc.paragraphs)
            if self._has_sect_pr(p._element)
        ]
        boundaries = []
        start = 0
        for b in breaks:
            boundaries.append((start, b))
            start = b + 1
        boundaries.append((start, len(doc.paragraphs) - 1))
        return boundaries

    @staticmethod
    def _has_sect_pr(p_elem) -> bool:
        pPr = p_elem.find(qn('w:pPr'))
        return pPr is not None and pPr.find(qn('w:sectPr')) is not None

    def _classify_section(self, doc: Document, start: int, end: int) -> str:
        """Classify section role from paragraph style names.
        Returns one of: 'content', 'toc', 'frontmatter', 'cover', 'structural_else'."""
        styles = set()
        for i in range(start, end + 1):
            s = (doc.paragraphs[i].style.name or "").lower()
            styles.add(s)

        if any("toc " in s for s in styles):
            return "toc"
        if "frontmatter" in styles:
            return "frontmatter"
        if "cap cover" in styles:
            return "cover"
        content_keywords = ("heading 1", "heading1", "body text",
                            "executive summary", "chapter number", "chapter title",
                            "appendix number")
        if any(kw in s for s in styles for kw in content_keywords):
            return "content"
        return "structural_else"

    def _slot_type(self, para) -> Optional[str]:
        """Return slot type for a template paragraph, or None if structural/navigational."""
        s = (para.style.name or "").lower()
        text = para.text.strip()

        # Cover title slots — only the first short text lines are title slots
        if "cover" in s:
            if text and len(text) < 150:
                return "heading"
            return None

        # Structural navigation / TOC
        if "toc " in s or "frontmatter" in s:
            return None

        # Numbered structural elements — these are design elements ("Chapter 1", "Appendix A")
        # and should NOT be filled with content text
        if "chapter number" in s or "appendix number" in s:
            return None

        # Headings
        for kw in ("heading 1", "heading1", "heading 2", "heading2",
                   "heading 3", "heading3", "heading 4", "heading4",
                   "executive summary/prelims heading"):
            if kw in s:
                return "heading"

        # Chapter title — this IS a content slot (it holds the chapter title text)
        if "chapter title" in s:
            return "heading"

        # Body text styles
        for kw in ("body text", "bodytext", "full out body text",
                   "body numbered prelims", "body annexes",
                   "bullets", "numbered bullets", "free numbering"):
            if kw in s:
                return "body"

        # Normal in a content section → likely body
        if s == "normal":
            if text and len(text) > 10:
                return "body"
            return None

        return None

    # ── Content classification ───────────────────────────

    def _pick_content_headings(self, paragraphs: List[ParagraphContent]) -> List[ParagraphContent]:
        """Filter paragraphs that map to heading slots (in original document order)."""
        result = []
        for p in paragraphs:
            role = self._resolve_content_role(p)
            if role in (SemanticRole.TITLE, SemanticRole.HEADING_1, SemanticRole.HEADING_2,
                        SemanticRole.HEADING_3, SemanticRole.HEADING_4):
                result.append(p)
        return result

    def _pick_content_body(self, paragraphs: List[ParagraphContent]) -> List[ParagraphContent]:
        """Filter paragraphs that map to body slots (in original document order)."""
        result = []
        for p in paragraphs:
            role = self._resolve_content_role(p)
            if role in (SemanticRole.BODY_TEXT, SemanticRole.LIST_BULLET,
                        SemanticRole.LIST_NUMBER, SemanticRole.UNKNOWN):
                result.append(p)
        return result

    # ── Replacement primitives ───────────────────────────

    @staticmethod
    def _replace_para_text_keep_structure(
        t_para_element,
        c_para: ParagraphContent,
        template: TemplateProfile,
        style_map: Dict[str, str],
    ) -> None:
        """Replace text in a template paragraph while keeping pPr (style, sectPr, etc.)."""
        for run in list(t_para_element.findall(qn('w:r'))):
            t_para_element.remove(run)

        if c_para.runs:
            for run_info in c_para.runs:
                new_run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                if run_info.get('bold'):
                    rPr.append(OxmlElement('w:b'))
                if run_info.get('italic'):
                    rPr.append(OxmlElement('w:i'))
                if run_info.get('underline'):
                    rPr.append(OxmlElement('w:u'))
                if len(rPr) > 0:
                    new_run.append(rPr)
                t_elem = OxmlElement('w:t')
                t_elem.text = run_info.get('text', '')
                new_run.append(t_elem)
                t_para_element.append(new_run)
        else:
            new_run = OxmlElement('w:r')
            t_elem = OxmlElement('w:t')
            t_elem.text = c_para.text or ''
            new_run.append(t_elem)
            t_para_element.append(new_run)

    @staticmethod
    def _clear_para_runs(p_elem) -> None:
        """Remove all runs from a paragraph (keeps pPr and sectPr intact)."""
        for run in list(p_elem.findall(qn('w:r'))):
            p_elem.remove(run)

    # ── Remaining content insertion ──────────────────────

    def _append_remaining_content(
        self,
        doc: Document,
        remaining: List[ParagraphContent],
        template: TemplateProfile,
        style_map: Dict[str, str],
    ) -> None:
        """Append remaining content paragraphs after the main body section."""
        if not remaining:
            return

        body_elem = doc.element.body
        insert_after = None
        # Find the last non-structural paragraph before the final sectPr
        for child in reversed(list(body_elem)):
            if child.tag == qn('w:p'):
                pPr = child.find(qn('w:pPr'))
                if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                    continue
                insert_after = child
                break
            elif child.tag == qn('w:tbl'):
                # Tables count too
                insert_after = child
                break

        for para in remaining:
            target_style_id = self._resolve_target_style_for_append(para, style_map, template)
            p_elem = self._new_paragraph_element(para, target_style_id)
            if insert_after is not None:
                idx = list(body_elem).index(insert_after)
                body_elem.insert(idx + 1, p_elem)
                insert_after = p_elem
            else:
                body_elem.append(p_elem)

    def _new_paragraph_element(self, para: ParagraphContent, style_id: Optional[str]):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        if style_id:
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), style_id)
            pPr.append(pStyle)
        p.append(pPr)

        if para.runs:
            for run_info in para.runs:
                r = OxmlElement('w:r')
                if run_info.get('bold') or run_info.get('italic') or run_info.get('underline'):
                    rPr = OxmlElement('w:rPr')
                    if run_info.get('bold'):
                        rPr.append(OxmlElement('w:b'))
                    if run_info.get('italic'):
                        rPr.append(OxmlElement('w:i'))
                    if run_info.get('underline'):
                        rPr.append(OxmlElement('w:u'))
                    r.append(rPr)
                t = OxmlElement('w:t')
                t.text = run_info.get('text', '')
                r.append(t)
                p.append(r)
        else:
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = para.text or ''
            r.append(t)
            p.append(r)
        return p

    def _resolve_target_style_for_append(
        self,
        para: ParagraphContent,
        style_map: Dict[str, str],
        template: TemplateProfile,
    ) -> Optional[str]:
        if para.style_name and para.style_name not in ('Normal', 'BodyText', 'BodyText2', 'NoSpacing'):
            if para.style_name in style_map:
                return style_map[para.style_name]
            if para.style_name in template.paragraph_styles:
                return para.style_name
        role = para.estimated_role
        if not role or role == SemanticRole.UNKNOWN:
            role = self._infer_role_from_content(para)
        if role == SemanticRole.HEADING_1:
            for sid, tstyle in template.paragraph_styles.items():
                if tstyle.semantic_role == SemanticRole.HEADING_1:
                    return sid
        return self._heuristic_style_match(para, template)

    # ── Legacy / fallback for fresh Document() creation ──

    def _build_fresh(
        self, doc: Document, template: TemplateProfile,
        content: ContentProfile, style_matches: List[StyleMatch],
    ) -> None:
        style_map = {m.source_style_id: m.target_style_id for m in style_matches}
        self._apply_document_defaults(doc, template)
        self._copy_template_styles(doc, template)
        for para in content.paragraphs:
            self._add_paragraph(doc, para, style_map, template)
        for table in content.tables:
            self._add_table(doc, table, style_map, template)
        self._copy_headers_footers(doc, template)

    def _apply_document_defaults(self, doc: Document, template: TemplateProfile) -> None:
        defaults = template.document_defaults
        if not defaults.margins:
            return
        try:
            section = doc.sections[0]
            if defaults.margins:
                for key in ('top', 'bottom', 'left', 'right'):
                    if key in defaults.margins:
                        setattr(section, f"{key}_margin", Inches(defaults.margins[key]))
            if defaults.page_width_inches:
                section.page_width = Inches(defaults.page_width_inches)
            if defaults.page_height_inches:
                section.page_height = Inches(defaults.page_height_inches)
            if defaults.orientation == 'landscape':
                section.orientation = 1
        except Exception as e:
            logger.warning(f"Could not apply document defaults: {e}")

    def _copy_template_styles(self, doc: Document, template: TemplateProfile) -> None:
        for style_id, tstyle in template.paragraph_styles.items():
            if tstyle.is_built_in and not tstyle.is_default:
                continue
            try:
                try:
                    existing = doc.styles[style_id]
                    self._apply_paragraph_style(existing, tstyle)
                except KeyError:
                    new_style = doc.styles.add_style(style_id, 1)
                    self._apply_paragraph_style(new_style, tstyle)
            except Exception as e:
                logger.debug(f"Could not copy style {style_id}: {e}")

    def _apply_paragraph_style(self, docx_style, tstyle: ParagraphStyle) -> None:
        try:
            font = docx_style.font
            if tstyle.font:
                if tstyle.font.name:
                    font.name = tstyle.font.name
                if tstyle.font.size_pt:
                    font.size = Pt(tstyle.font.size_pt)
                if tstyle.font.bold is not None:
                    font.bold = tstyle.font.bold
                if tstyle.font.italic is not None:
                    font.italic = tstyle.font.italic
                if tstyle.font.color and not tstyle.font.color.startswith('theme:'):
                    try:
                        font.color.rgb = RGBColor.from_string(tstyle.font.color)
                    except:
                        pass
            pfmt = docx_style.paragraph_format
            if tstyle.alignment and tstyle.alignment.alignment:
                align_map = {
                    '0': WD_ALIGN_PARAGRAPH.LEFT, 'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
                    '1': WD_ALIGN_PARAGRAPH.CENTER, 'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
                    '2': WD_ALIGN_PARAGRAPH.RIGHT, 'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
                    '3': WD_ALIGN_PARAGRAPH.JUSTIFY, 'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
                    'BOTH': WD_ALIGN_PARAGRAPH.JUSTIFY, 'None': WD_ALIGN_PARAGRAPH.LEFT,
                }
                if tstyle.alignment.alignment in align_map:
                    pfmt.alignment = align_map[tstyle.alignment.alignment]
            if tstyle.spacing:
                if tstyle.spacing.before_pt:
                    pfmt.space_before = Pt(tstyle.spacing.before_pt)
                if tstyle.spacing.after_pt:
                    pfmt.space_after = Pt(tstyle.spacing.after_pt)
                if tstyle.spacing.line_spacing:
                    pfmt.line_spacing = tstyle.spacing.line_spacing
            if tstyle.indentation:
                if tstyle.indentation.left_inches:
                    pfmt.left_indent = Inches(tstyle.indentation.left_inches)
                if tstyle.indentation.right_inches:
                    pfmt.right_indent = Inches(tstyle.indentation.right_inches)
                if tstyle.indentation.first_line_inches:
                    pfmt.first_line_indent = Inches(tstyle.indentation.first_line_inches)
        except Exception as e:
            logger.debug(f"Error applying style {tstyle.style_id}: {e}")

    def _add_paragraph(
        self, doc: Document, para: ParagraphContent,
        style_map: Dict[str, str], template: TemplateProfile,
    ) -> None:
        target_style_id = self._resolve_target_style(para, style_map, template)
        new_para = doc.add_paragraph()
        if para.runs:
            for run_info in para.runs:
                run = new_para.add_run(run_info.get('text', ''))
                self._apply_run_formatting(run, run_info)
        else:
            new_para.add_run(para.text)
        if target_style_id:
            try:
                new_para.style = target_style_id
            except KeyError:
                self._apply_style_by_definition(new_para, target_style_id, template)
        self._apply_direct_paragraph_formatting(new_para, para, template, target_style_id)

    def _resolve_target_style(
        self, para: ParagraphContent,
        style_map: Dict[str, str], template: TemplateProfile,
    ) -> Optional[str]:
        if para.style_name and para.style_name not in ('Normal', 'BodyText', 'BodyText2', 'NoSpacing'):
            if para.style_name in style_map:
                return style_map[para.style_name]
            if para.style_name in template.paragraph_styles:
                return para.style_name
        role = para.estimated_role
        if not role or role == SemanticRole.UNKNOWN:
            role = self._infer_role_from_content(para)
        if role and role != SemanticRole.UNKNOWN:
            for style_id, tstyle in template.paragraph_styles.items():
                if tstyle.semantic_role == role:
                    return style_id
            if role.value.startswith('heading') or role.value.startswith('title'):
                for style_id, tstyle in template.paragraph_styles.items():
                    trole = tstyle.semantic_role
                    if trole and (trole.value.startswith('heading') or trole.value.startswith('title')):
                        return style_id
        return self._heuristic_style_match(para, template)

    def _infer_role_from_content(self, para: ParagraphContent) -> SemanticRole:
        text = para.text.strip()
        if not text:
            return SemanticRole.BODY_TEXT
        if len(text) < 100:
            if para.max_font_size and para.max_font_size >= 20:
                return SemanticRole.TITLE
            if para.max_font_size and para.max_font_size >= 14:
                return SemanticRole.HEADING_1
            if para.has_bold:
                return SemanticRole.HEADING_2
        if para.is_list_item or text.startswith(('•', '●', '○', '▪', '-')):
            return SemanticRole.LIST_BULLET
        return SemanticRole.BODY_TEXT

    def _heuristic_style_match(self, para: ParagraphContent, template: TemplateProfile) -> Optional[str]:
        estimated_size = para.max_font_size or 11
        best_target = None
        best_score = 0.0
        for style_id, tstyle in template.paragraph_styles.items():
            score = 0.0
            checks = 0
            if tstyle.font and tstyle.font.size_pt:
                size_diff = abs(tstyle.font.size_pt - estimated_size)
                checks += 1
                if size_diff < 1:
                    score += 1.0
                elif size_diff < 3:
                    score += 0.7
                elif size_diff < 5:
                    score += 0.4
            if para.has_bold is not None and tstyle.font and tstyle.font.bold is not None:
                checks += 1
                value = 1.0 if para.has_bold == tstyle.font.bold else 0.0
                score += value
            if checks > 0 and score / checks > best_score:
                best_score = score / checks
                best_target = style_id
        return best_target if best_score >= 0.3 else None

    def _apply_style_by_definition(self, new_para, style_id: str, template: TemplateProfile) -> None:
        if style_id not in template.paragraph_styles:
            return
        tstyle = template.paragraph_styles[style_id]
        try:
            if tstyle.font:
                for run in new_para.runs:
                    if tstyle.font.name:
                        run.font.name = tstyle.font.name
                    if tstyle.font.size_pt:
                        run.font.size = Pt(tstyle.font.size_pt)
                    if tstyle.font.bold is not None:
                        run.font.bold = tstyle.font.bold
                    if tstyle.font.italic is not None:
                        run.font.italic = tstyle.font.italic
            pfmt = new_para.paragraph_format
            if tstyle.alignment and tstyle.alignment.alignment:
                align_map = {
                    '0': WD_ALIGN_PARAGRAPH.LEFT, 'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
                    '1': WD_ALIGN_PARAGRAPH.CENTER, 'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
                    '2': WD_ALIGN_PARAGRAPH.RIGHT, 'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
                    '3': WD_ALIGN_PARAGRAPH.JUSTIFY, 'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
                    'BOTH': WD_ALIGN_PARAGRAPH.JUSTIFY, 'None': WD_ALIGN_PARAGRAPH.LEFT,
                }
                if tstyle.alignment.alignment in align_map:
                    pfmt.alignment = align_map[tstyle.alignment.alignment]
            if tstyle.spacing:
                if tstyle.spacing.before_pt:
                    pfmt.space_before = Pt(tstyle.spacing.before_pt)
                if tstyle.spacing.after_pt:
                    pfmt.space_after = Pt(tstyle.spacing.after_pt)
                if tstyle.spacing.line_spacing:
                    pfmt.line_spacing = tstyle.spacing.line_spacing
            if tstyle.indentation:
                if tstyle.indentation.left_inches:
                    pfmt.left_indent = Inches(tstyle.indentation.left_inches)
                if tstyle.indentation.right_inches:
                    pfmt.right_indent = Inches(tstyle.indentation.right_inches)
                if tstyle.indentation.first_line_inches:
                    pfmt.first_line_indent = Inches(tstyle.indentation.first_line_inches)
        except Exception as e:
            logger.debug(f"Could not apply style definition {style_id}: {e}")

    def _apply_run_formatting(self, run, run_info: Dict[str, Any]) -> None:
        try:
            if run_info.get('bold'):
                run.bold = True
            if run_info.get('italic'):
                run.italic = True
            if run_info.get('underline'):
                run.underline = True
        except Exception:
            pass

    def _apply_direct_paragraph_formatting(
        self, new_para, para: ParagraphContent,
        template: TemplateProfile, target_style_id: Optional[str],
    ) -> None:
        try:
            pfmt = new_para.paragraph_format
            if para.is_list_item:
                indent = para.indentation_left_inches or 0.5 * para.list_level
                if indent > 0 and not pfmt.left_indent:
                    pfmt.left_indent = Inches(indent)
            if para.alignment and not target_style_id:
                align_map = {
                    'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
                    'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
                    'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
                    'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
                }
                if para.alignment in align_map:
                    pfmt.alignment = align_map[para.alignment]
        except Exception:
            pass

    def _add_table(self, doc: Document, table_info: Any,
                   style_map: Dict[str, str], template: TemplateProfile) -> None:
        row_count = getattr(table_info, 'row_count', 0) or 2
        col_count = getattr(table_info, 'column_count', 0) or 2
        try:
            table = doc.add_table(rows=row_count, cols=col_count)
            table.style = 'Table Grid'
            rows = getattr(table_info, 'rows', None)
            if rows:
                for i, row_data in enumerate(rows):
                    if i >= len(table.rows):
                        break
                    row = table.rows[i]
                    if isinstance(row_data, list):
                        for j, cell_text in enumerate(row_data):
                            if j >= len(row.cells):
                                break
                            row.cells[j].text = str(cell_text)
        except Exception as e:
            logger.warning(f"Could not add table: {e}")

    def _copy_headers_footers(self, doc: Document, template: TemplateProfile) -> None:
        try:
            if template.headers:
                section = doc.sections[0]
                if section.header is not None and template.headers.get('default'):
                    section.header.paragraphs[0].text = template.headers.get('default', '')
            if template.footers:
                section = doc.sections[0]
                if section.footer is not None and template.footers.get('default'):
                    section.footer.paragraphs[0].text = template.footers.get('default', '')
        except Exception as e:
            logger.warning(f"Could not copy headers/footers: {e}")
